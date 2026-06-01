import os
import io
from datetime import datetime
from backend.db import get_connection

SEVERITY_COLORS_RGB = {
    'Critical': (139, 0,   0),
    'High':     (233, 69,  96),
    'Medium':   (255, 140, 0),
    'Low':      (180, 150, 0),
    'Info':     (74,  158, 255),
}

SEVERITY_CVSS = {
    'Critical': '9.0 - 10.0',
    'High':     '7.0 - 8.9',
    'Medium':   '4.0 - 6.9',
    'Low':      '0.1 - 3.9',
    'Info':     'N/A',
}

BUSINESS_IMPACT = {
    'Critical': (
        'May lead to complete system compromise, unauthorised '
        'remote access, or full data breach.'
    ),
    'High': (
        'Could allow significant unauthorised access or data '
        'exposure with moderate exploitation effort.'
    ),
    'Medium': (
        'May expose sensitive information or enable limited '
        'unauthorised actions under specific conditions.'
    ),
    'Low': (
        'Minimal direct impact; contributes to the overall '
        'attack surface and reconnaissance value.'
    ),
    'Info': (
        'Informational only. No direct risk but may aid '
        'further enumeration by an attacker.'
    ),
}

REMEDIATION_EFFORT = {
    'Critical': 'Immediate',
    'High':     'Short-term',
    'Medium':   'Medium-term',
    'Low':      'Long-term',
    'Info':     'Optional',
}

TOOLS_INFO = [
    ('Nmap',         'Port scanning and service version detection'),
    ('Subfinder',    'Passive subdomain enumeration via OSINT'),
    ('httpx',        'HTTP probing to identify live web hosts'),
    ('WhatWeb',      'Web technology fingerprinting'),
    ('ffuf',         'Fast directory and endpoint fuzzing'),
    ('Nikto',        'Web server vulnerability scanning'),
    ('theHarvester', 'OSINT email and host harvesting'),
    ('DNSrecon',     'DNS enumeration and zone analysis'),
    ('Gobuster',     'Directory and file brute forcing'),
    ('Dirsearch',    'Web path discovery with curated wordlists'),
    ('WPScan',       'WordPress vulnerability scanning'),
    ('Nuclei',       'Template-based CVE vulnerability detection'),
]

ATTACK_PATHS = {
    'telnet':        'Attacker -> Network -> Port 23 -> Plaintext Sniff -> Full Access',
    'bindshell':     'Attacker -> Network -> Port 1524 -> Root Shell -> Full Compromise',
    'ftp':           'Attacker -> Network -> Port 21 -> Brute Force -> File Access',
    'mysql':         'Attacker -> Network -> Port 3306 -> Weak Credentials -> DB Access',
    'postgresql':    'Attacker -> Network -> Port 5432 -> Empty Password -> DB Access',
    'vnc':           'Attacker -> Network -> Port 5900 -> Default Login -> Desktop Control',
    'phpmyadmin':    'Attacker -> /phpMyAdmin -> Brute Force -> DB Admin -> RCE',
    'phpinfo':       'Attacker -> phpinfo.php -> Config Leak -> Targeted Exploit',
    'smb':           'Attacker -> Port 445 -> SMB Exploit -> Lateral Movement',
    'ssh':           'Attacker -> Port 22 -> Brute Force -> Remote Shell',
    'htpasswd':      'Attacker -> /.htpasswd -> Hash Download -> Crack -> Auth Bypass',
    'cve-2020-1938': 'Attacker -> AJP Port -> Ghostcat Exploit -> File Read -> RCE',
    'cve-2011-2523': 'Attacker -> vsftpd -> Backdoor Trigger -> Root Shell',
    'directory':     'Attacker -> Directory Listing -> File Enum -> Sensitive Data',
    'irc':           'Attacker -> IRC Port -> Version Detect -> Exploit -> RCE',
    'rmi':           'Attacker -> RMI Port -> Deserialization -> RCE',
    'smtp':          'Attacker -> Port 25 -> User Enum -> Phishing or Relay',
    'nfs':           'Attacker -> NFS Port -> Mount Share -> File Access',
    'wordpress':     'Attacker -> WPScan -> Plugin Vuln -> Admin Access -> RCE',
    'apache':        'Attacker -> Apache Version -> CVE Search -> Exploit -> Compromise',
    'htaccess':      'Attacker -> .htaccess -> Config Read -> Bypass -> Restricted Access',
}


def clean(text):
    if not text:
        return ''
    replacements = {
        '\u2014': '-',   '\u2013': '-',   '\u2019': "'",
        '\u2018': "'",   '\u201c': '"',   '\u201d': '"',
        '\u2022': '-',   '\u00e9': 'e',   '\u00e0': 'a',
        '\u00fc': 'u',   '\u00f6': 'o',   '\u00e4': 'a',
        '\u00b7': '-',   '\u2026': '...',  '\u00a0': ' ',
        '\u2012': '-',   '\u2015': '-',   '\u2010': '-',
        '\u2011': '-',   '\u00c2': '',    '\u00e2': 'a',
        '\u00c3': '',    '\u00e3': 'a',   '\u00b0': 'deg',
        '\u00ae': '(R)', '\u00a9': '(C)', '\u2122': '(TM)',
        '\u2192': '->',  '\u2190': '<-',  '\u2194': '<->',
    }
    result = str(text)
    for k, v in replacements.items():
        result = result.replace(k, v)
    result = result.encode('latin-1', errors='replace').decode('latin-1')
    return result


def truncate(text, max_len):
    text = clean(str(text))
    if len(text) > max_len:
        return text[:max_len - 3] + '...'
    return text


def get_scan_data(scan_id):
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM scans WHERE id=?', (scan_id,))
    scan = dict(cursor.fetchone())
    cursor.execute('''
        SELECT * FROM findings WHERE scan_id=?
        ORDER BY CASE severity
            WHEN "Critical" THEN 0
            WHEN "High"     THEN 1
            WHEN "Medium"   THEN 2
            WHEN "Low"      THEN 3
            WHEN "Info"     THEN 4
            ELSE 5 END
    ''', (scan_id,))
    findings = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return scan, findings


def count_by_severity(findings):
    counts = {
        'Critical': 0, 'High': 0,
        'Medium': 0, 'Low': 0, 'Info': 0
    }
    for f in findings:
        sev = f.get('severity', 'Info')
        counts[sev] = counts.get(sev, 0) + 1
    return counts


def get_finding_id(severity, index):
    prefix = {
        'Critical': 'AR-CRIT',
        'High':     'AR-HIGH',
        'Medium':   'AR-MED',
        'Low':      'AR-LOW',
        'Info':     'AR-INFO',
    }.get(severity, 'AR-FIND')
    return f"{prefix}-{index:03d}"


def get_attack_path(title, asset):
    combined = (title + ' ' + (asset or '')).lower()
    for keyword, path in ATTACK_PATHS.items():
        if keyword in combined:
            return path
    return None


def build_pie_chart(counts):
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    labels, sizes, colors = [], [], []
    color_map = {
        'Critical': '#8b0000',
        'High':     '#e94560',
        'Medium':   '#ff8c00',
        'Low':      '#b49600',
        'Info':     '#4a9eff',
    }
    for sev in ['Critical', 'High', 'Medium', 'Low', 'Info']:
        if counts.get(sev, 0) > 0:
            labels.append(f"{sev} ({counts[sev]})")
            sizes.append(counts[sev])
            colors.append(color_map[sev])

    fig, ax = plt.subplots(figsize=(5, 4), facecolor='white')
    ax.set_facecolor('white')
    if sizes:
        wedges, texts, autotexts = ax.pie(
            sizes, labels=labels, colors=colors,
            autopct='%1.1f%%', startangle=90,
            textprops={'color': '#333333', 'fontsize': 9}
        )
        for at in autotexts:
            at.set_color('#333333')
            at.set_fontsize(8)
    ax.set_title(
        'Findings by Severity',
        color='#333333', fontsize=11, pad=10,
        fontfamily='DejaVu Serif'
    )
    buf = io.BytesIO()
    fig.savefig(
        buf, format='png', dpi=130,
        bbox_inches='tight', facecolor='white'
    )
    plt.close(fig)
    buf.seek(0)
    return buf


def build_bar_chart(findings):
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    tool_counts = {}
    for f in findings:
        t = f['tool']
        tool_counts[t] = tool_counts.get(t, 0) + 1

    tools  = list(tool_counts.keys())
    counts = list(tool_counts.values())
    max_c  = max(counts) if counts else 1
    colors = [
        '#8b0000' if c == max_c else '#4a9eff'
        for c in counts
    ]

    fig, ax = plt.subplots(figsize=(8, 4), facecolor='white')
    ax.set_facecolor('#f8f8f8')
    if tools:
        bars = ax.barh(tools, counts, color=colors, height=0.55)
        ax.set_xlabel(
            'Number of Findings', color='#333333', fontsize=9
        )
        ax.tick_params(colors='#333333', labelsize=9)
        for spine in ax.spines.values():
            spine.set_color('#cccccc')
        ax.xaxis.grid(
            True, color='#dddddd', linestyle='--', alpha=0.7
        )
        ax.set_axisbelow(True)
        for bar, count in zip(bars, counts):
            ax.text(
                count + 0.1,
                bar.get_y() + bar.get_height() / 2,
                str(count),
                va='center', color='#333333', fontsize=9
            )
    ax.set_title(
        'Findings by Tool', color='#333333',
        fontsize=11, pad=10, fontfamily='DejaVu Serif'
    )
    buf = io.BytesIO()
    fig.savefig(
        buf, format='png', dpi=130,
        bbox_inches='tight', facecolor='white'
    )
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_executive_summary(scan, findings, counts):
    target   = clean(scan.get('target', 'the target'))
    total    = len(findings)
    critical = counts.get('Critical', 0)
    high     = counts.get('High', 0)
    medium   = counts.get('Medium', 0)

    para = (
        f"AutoRed conducted an automated reconnaissance and "
        f"vulnerability assessment against {target}. The assessment "
        f"identified a total of {total} findings across multiple "
        f"severity levels. "
    )
    if critical > 0:
        para += (
            f"Of these, {critical} finding(s) were rated Critical, "
            f"indicating immediately exploitable vulnerabilities "
            f"that pose a direct risk of system compromise or "
            f"data breach. "
        )
    if high > 0:
        para += (
            f"Additionally, {high} High-severity finding(s) were "
            f"identified which could be exploited with moderate "
            f"effort. "
        )
    if medium > 0:
        para += (
            f"{medium} Medium-severity finding(s) were also "
            f"detected, requiring attention in the near term. "
        )
    para += (
        f"Immediate remediation of Critical and High findings is "
        f"strongly recommended to reduce the overall attack surface."
    )
    return clean(para)


def generate_pdf(scan_id, output_path):
    from fpdf import FPDF

    scan, findings = get_scan_data(scan_id)
    counts = count_by_severity(findings)
    now    = datetime.now().strftime('%d %B %Y %H:%M')

    class PDF(FPDF):
        def header(self):
            if self.page_no() == 1:
                return
            self.set_font('Times', 'I', 10)
            self.set_text_color(150, 150, 150)
            self.cell(
                0, 8,
                'AutoRed Security Assessment Report  |  CONFIDENTIAL',
                ln=False, align='L'
            )
            self.cell(
                0, 8,
                f'Page {self.page_no()}',
                ln=True, align='R'
            )
            self.set_draw_color(200, 200, 200)
            self.line(15, self.get_y(), 195, self.get_y())
            self.ln(3)

        def footer(self):
            if self.page_no() == 1:
                return
            self.set_y(-15)
            self.set_font('Times', 'I', 10)
            self.set_text_color(150, 150, 150)
            self.set_draw_color(200, 200, 200)
            self.line(15, self.get_y(), 195, self.get_y())
            self.ln(1)
            self.cell(
                0, 6,
                f'CONFIDENTIAL  |  AutoRed v1.0  |  {now}',
                align='C'
            )

        def section_title(self, text, num=None):
            self.set_fill_color(233, 69, 96)
            self.rect(15, self.get_y(), 3, 12, 'F')
            self.set_font('Times', 'B', 14)
            self.set_text_color(233, 69, 96)
            prefix = f"{num}.  " if num else ""
            self.set_x(21)
            self.cell(
                0, 12,
                clean(prefix + text),
                ln=True
            )
            self.set_draw_color(230, 230, 230)
            self.line(15, self.get_y(), 195, self.get_y())
            self.ln(5)
            self.set_text_color(0, 0, 0)

        def subsection(self, text):
            self.set_font('Times', 'B', 13)
            self.set_text_color(50, 50, 120)
            self.cell(0, 8, clean(text), ln=True)
            self.set_text_color(30, 30, 30)

        def body_text(self, text):
            self.set_font('Times', '', 12)
            self.set_text_color(40, 40, 40)
            self.set_x(15)
            self.multi_cell(0, 7, clean(text), align='J')

        def kv_row(self, key, value,
                   key_color=(30, 80, 160),
                   val_color=(30, 30, 30)):
            self.set_font('Times', 'B', 12)
            self.set_text_color(*key_color)
            self.cell(50, 8, clean(key) + ':', ln=False)
            self.set_font('Times', '', 12)
            self.set_text_color(*val_color)
            val_text = truncate(value, 80)
            self.cell(0, 8, val_text, ln=True)

        def field_label(self, text):
            self.set_font('Times', 'B', 12)
            self.set_text_color(30, 80, 160)
            self.cell(0, 7, clean(text), ln=True)

        def field_value(self, text, max_len=300):
            self.set_font('Times', '', 12)
            self.set_text_color(50, 50, 50)
            val = truncate(text, max_len)
            self.multi_cell(0, 7, val)
            self.ln(1)

        def severity_badge(self, severity):
            color = SEVERITY_COLORS_RGB.get(
                severity, (100, 100, 100)
            )
            self.set_fill_color(*color)
            self.set_text_color(255, 255, 255)
            self.set_font('Times', 'B', 10)
            self.cell(
                30, 7, severity.upper(),
                border=0, fill=True,
                align='C', ln=False
            )
            self.set_text_color(30, 30, 30)

        def toc_entry(self, num, title, page):
            self.set_font('Times', '', 12)
            self.set_text_color(40, 40, 40)
            self.cell(10, 8, str(num) + '.', ln=False)
            self.cell(140, 8, clean(title), ln=False)
            self.set_font('Times', 'B', 12)
            self.set_text_color(233, 69, 96)
            self.cell(0, 8, str(page), ln=True, align='R')
            self.set_draw_color(230, 230, 230)
            self.line(15, self.get_y(), 195, self.get_y())

    pdf = PDF()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=True, margin=22)

    # ── COVER PAGE ──────────────────────────────────────────
    pdf.add_page()
    pdf.set_fill_color(13, 17, 23)
    pdf.rect(0, 0, 210, 297, 'F')

    pdf.set_fill_color(233, 69, 96)
    pdf.rect(0, 0, 210, 4, 'F')
    pdf.rect(0, 293, 210, 4, 'F')

    pdf.set_font('Times', 'B', 40)
    pdf.set_text_color(233, 69, 96)
    pdf.ln(30)
    pdf.cell(0, 18, 'AutoRed', ln=True, align='C')

    pdf.set_font('Times', '', 14)
    pdf.set_text_color(180, 188, 198)
    pdf.cell(
        0, 8,
        'Reconnaissance & Attack Surface Assessment Report',
        ln=True, align='C'
    )
    pdf.ln(6)

    pdf.set_draw_color(233, 69, 96)
    pdf.set_line_width(0.8)
    pdf.line(50, pdf.get_y(), 160, pdf.get_y())
    pdf.ln(12)

    pdf.set_fill_color(22, 27, 34)
    pdf.set_draw_color(48, 54, 61)
    pdf.set_line_width(0.3)
    box_y = pdf.get_y()
    pdf.rect(25, box_y, 160, 64, 'FD')
    pdf.ln(6)

    details = [
        ('Target',         truncate(scan.get('target', 'N/A'), 55)),
        ('Scan Name',      truncate(scan.get('name', 'N/A'), 55)),
        ('Profile',        truncate(scan.get('profile', 'N/A'), 40)),
        ('Date Generated', clean(now)),
        ('Prepared By',    'AutoRed v1.0'),
        ('Classification', 'CONFIDENTIAL'),
    ]
    for lbl, value in details:
        pdf.set_x(35)
        pdf.set_font('Times', 'B', 11)
        pdf.set_text_color(139, 148, 158)
        pdf.cell(52, 8, clean(lbl) + ':', ln=False)
        pdf.set_font('Times', '', 11)
        pdf.set_text_color(230, 237, 243)
        pdf.cell(0, 8, clean(str(value)), ln=True)

    pdf.ln(12)

    sev_order = ['Critical', 'High', 'Medium', 'Low', 'Info']
    col_w     = 28
    start_x   = (210 - col_w * 5) / 2

    pdf.set_font('Times', 'B', 11)
    pdf.set_text_color(139, 148, 158)
    pdf.cell(0, 8, 'FINDINGS OVERVIEW', ln=True, align='C')
    pdf.ln(3)

    pdf.set_x(start_x)
    for sev in sev_order:
        color = SEVERITY_COLORS_RGB.get(sev, (100, 100, 100))
        pdf.set_fill_color(*color)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font('Times', 'B', 16)
        pdf.cell(
            col_w, 14,
            str(counts.get(sev, 0)),
            border=0, fill=True,
            align='C', ln=False
        )
    pdf.ln()
    pdf.set_x(start_x)
    for sev in sev_order:
        pdf.set_font('Times', '', 9)
        pdf.set_text_color(139, 148, 158)
        pdf.cell(col_w, 6, sev, align='C', ln=False)
    pdf.ln(16)

    pdf.set_font('Times', 'I', 10)
    pdf.set_text_color(70, 80, 90)
    pdf.cell(
        0, 6,
        'CONFIDENTIAL - This document contains sensitive security information.',
        ln=True, align='C'
    )

    # ── TABLE OF CONTENTS ───────────────────────────────────
    pdf.add_page()
    pdf.section_title('Table of Contents')

    toc_items = [
        (1, 'Executive Summary',       3),
        (2, 'Severity Classification', 3),
        (3, 'Scope & Methodology',     4),
        (4, 'Findings Summary',        5),
        (5, 'Detailed Findings',       5),
        (6, 'Recommendations',         5 + max(1, len(findings) // 4)),
        (7, 'Conclusion',
         6 + max(1, len(findings) // 4)),
    ]
    for num, title, page in toc_items:
        pdf.toc_entry(num, title, page)
    pdf.ln(6)

    pdf.set_font('Times', 'I', 11)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(
        0, 7,
        clean(
            f'Total Findings: {len(findings)}  |  '
            f'Target: {scan.get("target","N/A")}  |  '
            f'Generated: {now}'
        ),
        ln=True
    )

    # ── EXECUTIVE SUMMARY ───────────────────────────────────
    pdf.add_page()
    pdf.section_title('Executive Summary', 1)
    pdf.body_text(
        generate_executive_summary(scan, findings, counts)
    )
    pdf.ln(8)

    # ── SEVERITY CLASSIFICATION ─────────────────────────────
    pdf.section_title('Severity Classification', 2)
    pdf.body_text(
        'All findings are classified according to standard '
        'vulnerability severity levels based on CVSS scoring '
        'guidelines. The table below defines each severity level, '
        'its CVSS range, business impact and recommended '
        'remediation timeline.'
    )
    pdf.ln(6)

    hdr_cols = ['Severity', 'CVSS Range', 'Business Impact', 'Action']
    hdr_ws   = [28, 28, 95, 28]
    pdf.set_fill_color(233, 69, 96)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font('Times', 'B', 11)
    for h, w in zip(hdr_cols, hdr_ws):
        pdf.cell(
            w, 9, h, border=1,
            fill=True, align='C', ln=False
        )
    pdf.ln()
    for sev in ['Critical', 'High', 'Medium', 'Low', 'Info']:
        color  = SEVERITY_COLORS_RGB.get(sev, (100, 100, 100))
        impact = truncate(BUSINESS_IMPACT[sev], 55)
        pdf.set_font('Times', 'B', 11)
        pdf.set_text_color(*color)
        pdf.cell(28, 8, sev, border=1, align='C', ln=False)
        pdf.set_text_color(40, 40, 40)
        pdf.set_font('Times', '', 11)
        pdf.cell(
            28, 8, SEVERITY_CVSS[sev],
            border=1, align='C', ln=False
        )
        pdf.cell(95, 8, impact, border=1, ln=False)
        pdf.cell(
            28, 8, REMEDIATION_EFFORT[sev],
            border=1, align='C', ln=True
        )

    # ── SCOPE & METHODOLOGY ─────────────────────────────────
    pdf.add_page()
    pdf.section_title('Scope & Methodology', 3)

    pdf.subsection('Scope of Assessment')
    pdf.ln(2)
    pdf.kv_row('Target',
               truncate(scan.get('target', 'N/A'), 70))
    pdf.kv_row('Scan Name',
               truncate(scan.get('name', 'N/A'), 70))
    pdf.kv_row('Profile',
               truncate(scan.get('profile', 'N/A'), 40))
    pdf.kv_row(
        'Approval',
        truncate(
            scan.get('approval_ref')
            or 'Authorized testing environment',
            70
        )
    )
    pdf.kv_row(
        'Conducted On',
        str(scan.get('created_at', 'N/A'))[:16]
    )
    pdf.ln(6)

    pdf.body_text(
        'This assessment was conducted within an authorized '
        'testing environment for educational and research purposes. '
        'All testing was performed using AutoRed v1.0, a GUI-based '
        'reconnaissance automation platform developed as a Final '
        'Year Project at Asia Pacific University (APU).'
    )
    pdf.ln(8)

    pdf.subsection('Assessment Phases')
    pdf.ln(2)
    phases = [
        ('1. Reconnaissance',
         'Passive and active information gathering using '
         'Subfinder, theHarvester and DNSrecon.'),
        ('2. Enumeration',
         'Port scanning and service detection via Nmap. '
         'Web host identification via httpx.'),
        ('3. Fingerprinting',
         'Technology stack identification using WhatWeb. '
         'CMS detection via WPScan.'),
        ('4. Vulnerability Discovery',
         'Web vulnerability scanning via Nikto and Nuclei. '
         'Directory enumeration via ffuf, Gobuster, Dirsearch.'),
        ('5. Classification',
         'All findings automatically parsed, normalized and '
         'severity-scored based on CVSS guidelines.'),
        ('6. Reporting',
         'Findings aggregated with business impact, '
         'attack paths and remediation guidance.'),
    ]
    for phase, desc in phases:
        pdf.set_font('Times', 'B', 12)
        pdf.set_text_color(233, 69, 96)
        pdf.cell(0, 7, clean(phase), ln=True)
        pdf.set_font('Times', '', 12)
        pdf.set_text_color(50, 50, 50)
        pdf.cell(0, 7, clean('    ' + desc), ln=True)
        pdf.ln(1)

    pdf.ln(4)
    pdf.subsection('Tools Used')
    pdf.ln(2)

    pdf.set_font('Times', 'B', 11)
    pdf.set_fill_color(233, 69, 96)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(
        38, 9, 'Tool',
        border=1, fill=True, align='C', ln=False
    )
    pdf.cell(
        137, 9, 'Purpose',
        border=1, fill=True, align='C', ln=True
    )
    for tool, desc in TOOLS_INFO:
        pdf.set_font('Times', 'B', 11)
        pdf.set_text_color(233, 69, 96)
        pdf.cell(38, 8, clean(tool), border=1, ln=False)
        pdf.set_font('Times', '', 11)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(137, 8, clean(desc), border=1, ln=True)

    # ── FINDINGS SUMMARY ────────────────────────────────────
    pdf.add_page()
    pdf.section_title('Findings Summary', 4)

    pdf.set_font('Times', 'B', 11)
    pdf.set_fill_color(233, 69, 96)
    pdf.set_text_color(255, 255, 255)
    sum_cols = ['Severity', 'Count', 'CVSS', 'Business Impact', 'Action']
    sum_ws   = [28, 15, 26, 88, 28]
    for c, w in zip(sum_cols, sum_ws):
        pdf.cell(
            w, 9, c, border=1,
            fill=True, align='C', ln=False
        )
    pdf.ln()
    for sev in ['Critical', 'High', 'Medium', 'Low', 'Info']:
        count  = counts.get(sev, 0)
        color  = SEVERITY_COLORS_RGB.get(sev, (100, 100, 100))
        impact = truncate(BUSINESS_IMPACT[sev], 56)
        pdf.set_font('Times', 'B', 11)
        pdf.set_text_color(*color)
        pdf.cell(28, 8, sev, border=1, align='C', ln=False)
        pdf.set_text_color(40, 40, 40)
        pdf.set_font('Times', '', 11)
        pdf.cell(
            15, 8, str(count),
            border=1, align='C', ln=False
        )
        pdf.cell(
            26, 8, SEVERITY_CVSS[sev],
            border=1, align='C', ln=False
        )
        pdf.cell(88, 8, impact, border=1, ln=False)
        pdf.cell(
            28, 8, REMEDIATION_EFFORT[sev],
            border=1, align='C', ln=True
        )

    pdf.ln(8)

    try:
        pie_buf = build_pie_chart(counts)
        bar_buf = build_bar_chart(findings)
        import tempfile
        with tempfile.NamedTemporaryFile(
            suffix='.png', delete=False
        ) as tf1:
            tf1.write(pie_buf.read())
            pie_path = tf1.name
        with tempfile.NamedTemporaryFile(
            suffix='.png', delete=False
        ) as tf2:
            tf2.write(bar_buf.read())
            bar_path = tf2.name

        pdf.set_font('Times', 'B', 12)
        pdf.set_text_color(30, 80, 160)
        pdf.cell(0, 7, 'Visual Analysis:', ln=True)
        pdf.ln(2)
        y_before = pdf.get_y()
        pdf.image(pie_path, x=15,  y=y_before, w=88, h=70)
        pdf.image(bar_path, x=110, y=y_before, w=85, h=70)
        pdf.ln(74)

        os.unlink(pie_path)
        os.unlink(bar_path)
    except Exception as e:
        print(f"[!] Chart generation skipped: {e}")

    # ── DETAILED FINDINGS ───────────────────────────────────
    pdf.add_page()
    pdf.section_title('Detailed Findings', 5)

    sev_counters = {
        'Critical': 0, 'High': 0,
        'Medium': 0, 'Low': 0, 'Info': 0
    }

    for finding in findings:
        if pdf.get_y() > 220:
            pdf.add_page()

        severity   = finding.get('severity', 'Info')
        sev_counters[severity] += 1
        finding_id = get_finding_id(
            severity, sev_counters[severity]
        )
        color = SEVERITY_COLORS_RGB.get(severity, (100, 100, 100))

        pdf.set_fill_color(245, 245, 245)
        pdf.set_draw_color(*color)
        pdf.set_line_width(0.6)
        block_y = pdf.get_y()
        pdf.rect(15, block_y, 180, 9, 'FD')

        pdf.set_font('Times', 'B', 11)
        pdf.set_text_color(*color)
        pdf.cell(38, 9, clean(finding_id), ln=False)
        pdf.set_text_color(20, 20, 20)
        title_text = truncate(finding.get('title', 'N/A'), 68)
        pdf.cell(0, 9, title_text, ln=True)

        pdf.set_line_width(0.2)
        pdf.set_draw_color(210, 210, 210)

        pdf.set_font('Times', 'B', 11)
        pdf.set_text_color(80, 80, 80)
        pdf.cell(24, 7, 'Severity:', ln=False)
        pdf.severity_badge(severity)
        pdf.cell(8, 7, '', ln=False)

        pdf.set_text_color(80, 80, 80)
        pdf.cell(16, 7, 'CVSS:', ln=False)
        pdf.set_font('Times', '', 11)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(26, 7, SEVERITY_CVSS[severity], ln=False)

        pdf.set_font('Times', 'B', 11)
        pdf.set_text_color(80, 80, 80)
        pdf.cell(14, 7, 'Tool:', ln=False)
        pdf.set_font('Times', '', 11)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(
            26, 7,
            truncate(finding.get('tool', 'N/A'), 18),
            ln=False
        )
        pdf.set_font('Times', 'B', 11)
        pdf.set_text_color(80, 80, 80)
        pdf.cell(16, 7, 'Action:', ln=False)
        pdf.set_font('Times', '', 11)
        pdf.set_text_color(*color)
        pdf.cell(
            0, 7,
            REMEDIATION_EFFORT.get(severity, 'Review'),
            ln=True
        )

        pdf.set_font('Times', 'B', 11)
        pdf.set_text_color(80, 80, 80)
        pdf.cell(16, 7, 'Asset:', ln=False)
        pdf.set_font('Times', '', 11)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(
            0, 7,
            truncate(finding.get('asset', 'N/A'), 85),
            ln=True
        )
        pdf.ln(2)

        desc = str(finding.get('description', '') or '')
        if desc:
            pdf.field_label('Description:')
            pdf.field_value(desc, 300)

        pdf.field_label('Business Impact:')
        pdf.field_value(BUSINESS_IMPACT.get(severity, ''), 200)

        attack_path = get_attack_path(
            finding.get('title', ''),
            finding.get('asset', '')
        )
        if attack_path:
            pdf.field_label('Attack Path:')
            pdf.set_font('Times', '', 12)
            pdf.set_text_color(180, 60, 60)
            pdf.cell(0, 6, clean(attack_path), ln=True)
            pdf.ln(1)

        rec = str(finding.get('recommendation', '') or '')
        if rec:
            pdf.field_label('Recommendation:')
            pdf.field_value(rec, 250)

        pdf.ln(3)
        pdf.set_draw_color(210, 210, 210)
        pdf.set_line_width(0.2)
        pdf.line(15, pdf.get_y(), 195, pdf.get_y())
        pdf.ln(5)

    # ── RECOMMENDATIONS ─────────────────────────────────────
    pdf.add_page()
    pdf.section_title('Recommendations', 6)
    pdf.body_text(
        'The following prioritised recommendations are based on '
        'the severity and exploitability of identified findings. '
        'Critical and High items should be addressed immediately '
        'to reduce the risk of exploitation.'
    )
    pdf.ln(6)

    priority = [
        f for f in findings
        if f.get('severity') in ['Critical', 'High']
        and f.get('recommendation')
    ][:10]

    rc = {
        'Critical': 0, 'High': 0,
        'Medium': 0, 'Low': 0, 'Info': 0
    }
    for f in priority:
        sev   = f['severity']
        rc[sev] += 1
        fid   = get_finding_id(sev, rc[sev])
        color = SEVERITY_COLORS_RGB.get(sev, (100, 100, 100))

        pdf.set_font('Times', 'B', 12)
        pdf.set_text_color(*color)
        pdf.cell(
            0, 8,
            clean(f'[{fid}]  {truncate(f["title"], 65)}'),
            ln=True
        )
        pdf.set_font('Times', '', 12)
        pdf.set_text_color(50, 50, 50)
        rec = truncate(f['recommendation'], 220)
        pdf.multi_cell(0, 7, rec)
        pdf.ln(4)

    # ── CONCLUSION ──────────────────────────────────────────
    pdf.add_page()
    pdf.section_title('Conclusion', 7)

    target = truncate(scan.get('target', 'the target'), 60)
    pdf.body_text(
        f"This automated reconnaissance assessment of {target} "
        f"identified a total of {len(findings)} findings across "
        f"five severity levels, including "
        f"{counts.get('Critical', 0)} Critical and "
        f"{counts.get('High', 0)} High severity findings that "
        f"require immediate attention."
    )
    pdf.ln(6)
    pdf.body_text(
        'Immediate action is recommended for all Critical and '
        'High findings, particularly any exposed remote access '
        'services, known backdoors, or unpatched vulnerabilities. '
        'Medium findings should be scheduled for remediation in '
        'the near term, while Low and Informational findings '
        'should be reviewed as part of ongoing security hygiene.'
    )
    pdf.ln(6)
    pdf.body_text(
        'AutoRed demonstrated significant time savings compared '
        'to manual reconnaissance, completing the assessment in '
        'under 10 minutes while identifying findings that would '
        'typically require hours of manual effort. This validates '
        'the platform as an effective tool for initial attack '
        'surface assessment in cybersecurity engagements.'
    )
    pdf.ln(10)

    pdf.set_font('Times', 'B', 13)
    pdf.set_text_color(233, 69, 96)
    pdf.cell(0, 8, 'Disclaimer', ln=True)
    pdf.set_draw_color(233, 69, 96)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(4)
    pdf.set_font('Times', 'I', 11)
    pdf.set_text_color(100, 100, 100)
    pdf.body_text(
        'This report was generated automatically by AutoRed v1.0 '
        'for authorized security assessment and educational '
        'purposes only. All testing was conducted within an '
        'approved scope. This report is classified as '
        'CONFIDENTIAL and should be handled accordingly.'
    )

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    pdf.output(output_path)
    print(f"[+] PDF report saved to: {output_path}")
    return output_path


def generate_docx(scan_id, output_path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    scan, findings = get_scan_data(scan_id)
    counts = count_by_severity(findings)
    now    = datetime.now().strftime('%d %B %Y %H:%M')

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    def add_heading(text, level=1, color=None):
        h = doc.add_heading(clean(text), level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if color and h.runs:
            for run in h.runs:
                run.font.color.rgb = RGBColor(*color)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14 if level == 1 else 13)
        return h

    def add_para(text, bold=False, color=None, size=12):
        p   = doc.add_paragraph()
        run = p.add_run(clean(text))
        run.bold = bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    def add_kv(key, value):
        p  = doc.add_paragraph()
        kr = p.add_run(clean(key) + ': ')
        kr.bold = True
        kr.font.name = 'Times New Roman'
        kr.font.size = Pt(12)
        kr.font.color.rgb = RGBColor(30, 80, 160)
        vr = p.add_run(truncate(str(value), 100))
        vr.font.name = 'Times New Roman'
        vr.font.size = Pt(12)
        vr.font.color.rgb = RGBColor(30, 30, 30)
        return p

    title_h = doc.add_heading('AutoRed', 0)
    title_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_h.runs:
        run.font.color.rgb = RGBColor(233, 69, 96)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(32)

    sub = doc.add_paragraph(
        'Reconnaissance & Attack Surface Assessment Report'
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].font.size = Pt(14)
        sub.runs[0].font.name = 'Times New Roman'
        sub.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    add_heading('Scan Details', 1, (233, 69, 96))
    details = [
        ('Target',         scan.get('target', 'N/A')),
        ('Scan Name',      scan.get('name', 'N/A')),
        ('Profile',        scan.get('profile', 'N/A')),
        ('Date Generated', now),
        ('Approval Ref',   scan.get('approval_ref') or 'N/A'),
        ('Classification', 'CONFIDENTIAL'),
        ('Total Findings', str(len(findings))),
    ]
    tbl = doc.add_table(rows=len(details), cols=2)
    tbl.style = 'Table Grid'
    for i, (k, v) in enumerate(details):
        tbl.cell(i, 0).text = clean(k)
        tbl.cell(i, 1).text = truncate(str(v), 80)
        for para in tbl.cell(i, 0).paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        for para in tbl.cell(i, 1).paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    doc.add_page_break()
    add_heading('Executive Summary', 1, (233, 69, 96))
    p = doc.add_paragraph(
        generate_executive_summary(scan, findings, counts)
    )
    if p.runs:
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(12)

    doc.add_page_break()
    add_heading('Findings Summary', 1, (233, 69, 96))
    sum_tbl = doc.add_table(rows=1, cols=4)
    sum_tbl.style = 'Table Grid'
    hdr = sum_tbl.rows[0].cells
    for i, h in enumerate(
        ['Severity', 'Count', 'CVSS', 'Action']
    ):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    for sev in ['Critical', 'High', 'Medium', 'Low', 'Info']:
        row = sum_tbl.add_row().cells
        row[0].text = sev
        row[1].text = str(counts.get(sev, 0))
        row[2].text = SEVERITY_CVSS[sev]
        row[3].text = REMEDIATION_EFFORT[sev]
        color = SEVERITY_COLORS_RGB.get(sev, (0, 0, 0))
        for para in row[0].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(*color)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    doc.add_page_break()
    add_heading('Detailed Findings', 1, (233, 69, 96))

    sev_counters = {
        'Critical': 0, 'High': 0,
        'Medium': 0, 'Low': 0, 'Info': 0
    }

    for finding in findings:
        sev = finding.get('severity', 'Info')
        sev_counters[sev] += 1
        fid   = get_finding_id(sev, sev_counters[sev])
        color = SEVERITY_COLORS_RGB.get(sev, (100, 100, 100))

        h = doc.add_heading(
            clean(
                f"{fid} - "
                f"{truncate(finding.get('title','N/A'), 70)}"
            ),
            level=2
        )
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)

        meta = doc.add_paragraph()
        for k, v in [
            ('Severity',  sev),
            ('CVSS',      SEVERITY_CVSS[sev]),
            ('Tool',      finding.get('tool', 'N/A')),
            ('Asset',     truncate(finding.get('asset','N/A'), 60)),
            ('Action',    REMEDIATION_EFFORT[sev]),
        ]:
            kr = meta.add_run(f'{k}: ')
            kr.bold = True
            kr.font.name = 'Times New Roman'
            kr.font.size = Pt(12)
            kr.font.color.rgb = RGBColor(30, 80, 160)
            vr = meta.add_run(clean(str(v)) + '   ')
            vr.font.name = 'Times New Roman'
            vr.font.size = Pt(12)
            vr.font.color.rgb = RGBColor(50, 50, 50)

        for section, key in [
            ('Description',    'description'),
            ('Business Impact', None),
            ('Attack Path',     None),
            ('Recommendation', 'recommendation'),
        ]:
            p  = doc.add_paragraph()
            kr = p.add_run(section + ': ')
            kr.bold = True
            kr.font.name = 'Times New Roman'
            kr.font.size = Pt(12)
            kr.font.color.rgb = RGBColor(30, 80, 160)

            if key:
                val = truncate(
                    str(finding.get(key, '') or ''), 300
                )
            elif section == 'Business Impact':
                val = clean(BUSINESS_IMPACT.get(sev, ''))
            else:
                ap = get_attack_path(
                    finding.get('title', ''),
                    finding.get('asset', '')
                )
                val = clean(ap) if ap \
                    else 'Not identified for this finding.'

            vr = p.add_run(val)
            vr.font.name = 'Times New Roman'
            vr.font.size = Pt(12)
            vr.font.color.rgb = RGBColor(50, 50, 50)

    doc.add_page_break()
    add_heading('Methodology', 1, (233, 69, 96))
    mp = doc.add_paragraph(
        'This assessment was conducted using AutoRed v1.0, '
        'a GUI-based reconnaissance automation platform developed '
        'as a Final Year Project at Asia Pacific University (APU). '
        'All testing was performed within an authorized environment.'
    )
    if mp.runs:
        mp.runs[0].font.name = 'Times New Roman'
        mp.runs[0].font.size = Pt(12)

    m_tbl = doc.add_table(rows=1, cols=2)
    m_tbl.style = 'Table Grid'
    m_tbl.rows[0].cells[0].text = 'Tool'
    m_tbl.rows[0].cells[1].text = 'Purpose'
    for para in m_tbl.rows[0].cells[0].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    for para in m_tbl.rows[0].cells[1].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    for tool, desc in TOOLS_INFO:
        row = m_tbl.add_row().cells
        row[0].text = clean(tool)
        row[1].text = clean(desc)

    doc.add_page_break()
    add_heading('Conclusion', 1, (233, 69, 96))
    cp = doc.add_paragraph(
        clean(
            f"This assessment of "
            f"{truncate(scan.get('target','the target'), 60)} "
            f"identified {len(findings)} findings. "
            f"Immediate remediation of Critical and High findings "
            f"is strongly recommended."
        )
    )
    if cp.runs:
        cp.runs[0].font.name = 'Times New Roman'
        cp.runs[0].font.size = Pt(12)

    add_heading('Disclaimer', 2, (100, 100, 100))
    disc = doc.add_paragraph(
        'This report was generated by AutoRed v1.0 for authorized '
        'security assessment purposes only. Treat as CONFIDENTIAL.'
    )
    if disc.runs:
        disc.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        disc.runs[0].font.italic = True
        disc.runs[0].font.name = 'Times New Roman'
        disc.runs[0].font.size = Pt(12)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"[+] DOCX report saved to: {output_path}")
    return output_path


if __name__ == '__main__':
    from backend.db import init_db
    init_db()
    scan_id   = 27
    pdf_path  = f'storage/{scan_id}/report/report.pdf'
    docx_path = f'storage/{scan_id}/report/report.docx'
    print("[*] Generating PDF...")
    generate_pdf(scan_id, pdf_path)
    print("[*] Generating DOCX...")
    generate_docx(scan_id, docx_path)
    print("[+] Done!")
